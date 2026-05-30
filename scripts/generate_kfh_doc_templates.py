from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "docs" / "document-templates"


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_padding(cell, top=90, right=90, bottom=90, left=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
      tc_mar = OxmlElement("w:tcMar")
      tc_pr.append(tc_mar)

    for margin_name, margin_value in (
        ("top", top),
        ("right", right),
        ("bottom", bottom),
        ("left", left),
    ):
        element = tc_mar.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D1D5DB")


def apply_base_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)


def add_header_meta(document: Document, template_kind: str) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(template_kind)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Шаблон для загрузки в AgroScope")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("Дата формирования: [document_date]")
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_kfh_block(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run("КФХ/организация: ").bold = True
    paragraph.add_run("______________________________")

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run("Ответственное лицо: ").bold = True
    paragraph.add_run("______________________________")


def add_signatures_block(document: Document, labels: list[str]) -> None:
    document.add_paragraph()
    table = document.add_table(rows=len(labels), cols=2)
    table.autofit = False
    set_table_borders(table)

    for row, label in zip(table.rows, labels):
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.0)
        row.cells[0].text = label
        row.cells[1].text = "______________________________"
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_padding(cell, top=120, right=100, bottom=120, left=100)


def style_table(table) -> None:
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_padding(cell)


def create_single_operation_act() -> Path:
    document = Document()
    apply_base_styles(document)
    add_header_meta(document, "Внутренний акт выполнения полевой операции")
    add_kfh_block(document)

    document.add_paragraph()
    intro = document.add_paragraph()
    intro.add_run(
        "Настоящим подтверждается выполнение полевой операции на основании фактических данных в системе AgroScope."
    )

    table = document.add_table(rows=0, cols=2)
    style_table(table)
    rows = [
        ("Период/дата операции", "[period_start]"),
        ("Номер операции", "[operation_id]"),
        ("Дата операции", "[operation_date]"),
        ("Тип операции", "[operation_type]"),
        ("Статус операции", "[operation_status]"),
        ("Календарный статус", "[operation_calendar_status]"),
        ("Поле", "[field_name]"),
        ("Идентификатор поля", "[field_id]"),
        ("Площадь поля, га", "[field_area]"),
        ("Кадастровый номер", "[field_cadastral_number]"),
        ("Культура", "[crop_name]"),
        ("Код культуры", "[crop_code]"),
        ("Время начала", "[time_start]"),
        ("Время окончания", "[time_end]"),
        ("Техника/агрегат", "[equipment_name]"),
        ("Тип техники", "[equipment_type]"),
        ("Модель техники", "[equipment_model]"),
        ("Госномер", "[equipment_reg_number]"),
    ]

    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].width = Cm(7.0)
        row.cells[1].width = Cm(9.8)

    document.add_paragraph()
    note = document.add_paragraph()
    note.add_run("Примечание: ").bold = True
    note.add_run(
        "Шаблон предназначен для внутреннего акта по одной операции. При необходимости хозяйство может адаптировать формулировки под свою учетную политику."
    )

    add_signatures_block(
        document,
        ["Исполнитель/агроном", "Проверил(а)", "Утвердил(а)"],
    )

    output = OUTPUT_DIR / "vnutrenniy-akt-vypolneniya-polevoy-operacii.docx"
    document.save(output)
    return output


def create_daily_registry() -> Path:
    document = Document()
    apply_base_styles(document)
    add_header_meta(document, "Ежедневный реестр выполненных полевых работ")
    add_kfh_block(document)

    filters = document.add_table(rows=2, cols=4)
    style_table(filters)
    filters.cell(0, 0).text = "Дата реестра"
    filters.cell(0, 1).text = "[period_start]"
    filters.cell(0, 2).text = "Количество операций"
    filters.cell(0, 3).text = "[operations_count]"
    filters.cell(1, 0).text = "Поле (если фильтр)"
    filters.cell(1, 1).text = "[field_name]"
    filters.cell(1, 2).text = "Площадь поля"
    filters.cell(1, 3).text = "[field_area]"

    document.add_paragraph()
    table = document.add_table(rows=2, cols=7)
    style_table(table)
    headers = [
        "№",
        "Дата",
        "Поле",
        "Тип операции",
        "Статус",
        "Культура",
        "Техника",
    ]
    for index, label in enumerate(headers):
        table.cell(0, index).text = label

    marker_row = table.rows[1].cells
    marker_row[0].text = "[items_start][item_index]"
    marker_row[1].text = "[item_operation_date]"
    marker_row[2].text = "[item_field_name]"
    marker_row[3].text = "[item_operation_type]"
    marker_row[4].text = "[item_operation_calendar_status]"
    marker_row[5].text = "[item_crop_name]"
    marker_row[6].text = "[item_equipment_name][items_end]"

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.add_run("Подтверждение: ").bold = True
    note.add_run(
        "реестр составлен по выполненным операциям за выбранный день и используется как внутренняя первичная сводка для агронома и бухгалтерии."
    )

    add_signatures_block(document, ["Составил(а)", "Проверил(а)"])

    output = OUTPUT_DIR / "ezhednevnyy-reestr-vypolnennykh-polevykh-rabot.docx"
    document.save(output)
    return output


def create_period_registry() -> Path:
    document = Document()
    apply_base_styles(document)
    add_header_meta(document, "Реестр выполненных полевых работ за период")
    add_kfh_block(document)

    filters = document.add_table(rows=2, cols=4)
    style_table(filters)
    filters.cell(0, 0).text = "Период с"
    filters.cell(0, 1).text = "[period_start]"
    filters.cell(0, 2).text = "Период по"
    filters.cell(0, 3).text = "[period_end]"
    filters.cell(1, 0).text = "Поле (если фильтр)"
    filters.cell(1, 1).text = "[field_name]"
    filters.cell(1, 2).text = "Количество операций"
    filters.cell(1, 3).text = "[operations_count]"

    document.add_paragraph()
    table = document.add_table(rows=2, cols=10)
    style_table(table)
    headers = [
        "№",
        "Дата",
        "Поле",
        "Площадь, га",
        "Тип операции",
        "Статус",
        "Культура",
        "Время начала",
        "Время окончания",
        "Техника",
    ]
    for index, label in enumerate(headers):
        table.cell(0, index).text = label

    marker_row = table.rows[1].cells
    marker_row[0].text = "[items_start][item_index]"
    marker_row[1].text = "[item_operation_date]"
    marker_row[2].text = "[item_field_name]"
    marker_row[3].text = "[item_field_area]"
    marker_row[4].text = "[item_operation_type]"
    marker_row[5].text = "[item_operation_calendar_status]"
    marker_row[6].text = "[item_crop_name]"
    marker_row[7].text = "[item_time_start]"
    marker_row[8].text = "[item_time_end]"
    marker_row[9].text = "[item_equipment_name][items_end]"

    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(10)
    summary.add_run("Итоговая отметка: ").bold = True
    summary.add_run(
        "реестр предназначен для внутреннего контроля выполненных полевых работ за период и может использоваться как основа для дальнейшего оформления бухгалтерских документов."
    )

    add_signatures_block(document, ["Составил(а)", "Проверил(а)", "Утвердил(а)"])

    output = OUTPUT_DIR / "reestr-vypolnennykh-rabot-za-period.docx"
    document.save(output)
    return output


def main() -> None:
    ensure_output_dir()
    outputs = [
        create_single_operation_act(),
        create_daily_registry(),
        create_period_registry(),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
